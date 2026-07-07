from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from scripts.build_github_project_dossier import (
    BLUE,
    DARK_BLUE,
    INK,
    LIGHT_BLUE,
    MUTED,
    PROJECTS,
    RESUME_VARIANTS,
    WORK_EXPERIENCE,
    add_label_paragraph,
    add_page_field,
    set_font,
    setup_styles,
)


ROOT = Path(__file__).resolve().parents[1]
PROFILE_DIR = ROOT / "profile"
CONTEXT_PATH = PROFILE_DIR / "candidate_context.json"
TEMPLATE_PATH = PROFILE_DIR / "resume_template.tex"
DOCX_PATH = ROOT / "artifacts" / "Candidate_Profile_Context_Tushar_Pathak.docx"
SOURCE_TEMPLATE = Path(
    "/Users/tushardhananjaypathak/.codex/attachments/02565ed8-6008-4de9-9aa4-12dcb43c9b96/pasted-text.txt"
)


IDENTITY = {
    "resume_name": "Tushar Pathak",
    "full_name_seen_in_records": "Tushar Dhananjay Pathak / Pathak Tushar Dhananjay",
    "phone": "+1 (718) 413-9793",
    "email": "tpathak1411@gmail.com",
    "linkedin": "https://linkedin.com/in/tushar-pathak-7b945219a/",
    "github": "https://github.com/tusharpathaknyu",
    "portfolio": "https://tusharpathaknyu.github.io/",
}


EDUCATION = [
    {
        "institution": "New York University",
        "location": "New York, USA",
        "degree": "M.S. Computer Engineering",
        "dates": "January 2024 - December 2025",
        "additional_context": ["Graduate Scholarship Recipient"],
    },
    {
        "institution": "Birla Institute of Technology and Science (BITS) Pilani",
        "location": "Hyderabad, India",
        "degree": "B.E. Electrical and Electronics Engineering",
        "dates": "August 2019 - July 2023",
        "additional_context": [],
    },
]


TARGET_LANES = [
    "Power electronics and power-conversion engineering",
    "Analog, board and hardware validation",
    "FPGA design, prototyping and validation engineering",
    "SoC, ASIC, RTL and verification engineering",
    "SerDes and mixed-signal validation",
    "AI/ML engineering for circuits and hardware",
    "Applications and field-applications engineering",
    "Solutions engineering and customer-facing technical roles",
    "Software engineering: backend, full-stack and platform roles",
    "Applied AI, machine-learning and AI product engineering",
    "Game development, Unity/Godot gameplay and XR/fitness-game engineering",
]


SKILL_GROUPS = {
    "programming_and_scripting": [
        "Python", "C", "C++", "SystemVerilog", "Verilog", "Tcl", "Bash", "MATLAB"
    ],
    "digital_design_and_verification": [
        "RTL design", "UVM", "SVA", "RISC-V RV32I", "cocotb", "Icarus Verilog",
        "Verilator", "Yosys", "Vivado", "Gowin EDA", "functional and code coverage"
    ],
    "power_and_mixed_signal": [
        "DC-DC converters", "SPICE/ngspice", "LTspice", "Altium Designer", "KiCad",
        "PCB design", "oscilloscope-based validation", "S-parameters", "SerDes equalization",
        "BER and eye analysis"
    ],
    "ai_ml_and_data": [
        "PyTorch", "reinforcement learning", "transformers", "graph models", "Gymnasium",
        "Stable-Baselines3", "NumPy", "Pandas", "Matplotlib", "Jupyter"
    ],
    "software_and_platforms": [
        "REST APIs", "FastAPI", "backend services", "Swagger/OpenAPI", "Streamlit",
        "Docker", "Cloud Run", "GitHub Actions", "Datadog"
    ],
    "game_and_interactive_systems": [
        "Unity 6", "C#", "Godot 4.5", "GDScript", "React Native Unity bridge",
        "MediaPipe Pose Landmarker", "camera-driven gameplay", "game progression systems",
        "biometric PvP mechanics", "mobile game integration", "XR/VR/AR-adjacent interaction"
    ],
}


def clean_experience() -> list[dict]:
    result = []
    for item in WORK_EXPERIENCE:
        result.append(
            {
                "organization": item["organization"],
                "role": item["role"],
                "dates": item["dates"],
                "location": item["location"],
                "responsibilities_and_contributions": item["safe"],
                "internal_source_context": item["evidence"],
                "internal_tailoring_note": item["verify"],
            }
        )
    return result


def clean_projects() -> list[dict]:
    result = []
    for project in PROJECTS:
        result.append(
            {
                "name": project["name"],
                "repository": project["repo"],
                "url": f"https://github.com/tusharpathaknyu/{project['repo']}",
                "category": project["category"],
                "summary": project["summary"],
                "highlights": project["highlights"],
                "technology_context": project["tech"],
                "canonical_resume_bullet": project["resume"],
                "internal_source_context": project["sources"],
                "internal_tailoring_note": project["verify"],
            }
        )
    return result


def build_context() -> dict:
    return {
        "schema_version": 1,
        "updated": date.today().isoformat(),
        "purpose": "Canonical candidate memory for job discovery, resume tailoring and approval-first applications.",
        "identity": IDENTITY,
        "professional_positioning": {
            "primary": "Electrical and computer engineer with experience spanning power electronics, FPGA/RTL, backend software development and applied AI.",
            "role_specific_rule": "Vary the emphasis among hardware, software and AI according to the job while retaining the candidate's electrical and computer engineering foundation.",
            "prohibited_headline": "Do not describe the candidate as a Software and AI Engineer.",
        },
        "education": EDUCATION,
        "experience": clean_experience(),
        "target_lanes": TARGET_LANES,
        "skill_groups": SKILL_GROUPS,
        "resume_style": {
            "format": "LaTeX",
            "template_file": str(TEMPLATE_PATH.relative_to(ROOT)),
            "page_size": "US Letter",
            "body_font": "Charter, 9 pt",
            "page_count_target": 1,
            "section_order": ["Education", "Work Experience", "Role-specific Projects", "Technical Skills"],
            "selection_rules": [
                "Preserve the supplied preamble, commands, contact block and section-rule styling.",
                "Select two or three projects that best match the job description.",
                "Reorder experience or projects for relevance without changing dates, titles or facts.",
                "Keep concise action-and-impact bullets and return complete compilable LaTeX.",
            ],
            "recent_resume_variants": [
                {
                    "timestamp": timestamp,
                    "target_lane": lane,
                    "experience_order": experience,
                    "selected_projects": projects,
                }
                for timestamp, lane, experience, projects, _note in RESUME_VARIANTS
            ],
        },
        "projects": clean_projects(),
        "profile_fields_to_collect_later": [
            "Preferred job locations and remote/hybrid constraints",
            "Work authorization and sponsorship needs",
            "Salary preferences",
            "Roles or industries to exclude",
            "Final approved wording for any new work or project results",
        ],
    }


def set_repeat_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("CANDIDATE PROFILE CONTEXT"), size=9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Agent context • "), size=9, color=MUTED)
    add_page_field(footer)
    for run in footer.runs:
        set_font(run, size=9, color=MUTED)


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(30)
    kicker.paragraph_format.space_after = Pt(10)
    set_font(kicker.add_run("JOB APPLICATION AGENT MEMORY"), size=11, bold=True, color=DARK_BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run("Candidate Profile Context"), size=28, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    set_font(subtitle.add_run("Tushar Pathak • Work, projects, skills and resume house style"), size=14, color=DARK_BLUE)

    add_label_paragraph(
        doc,
        "Purpose",
        "A readable copy of the information the application agent can use when finding roles and tailoring resumes. This is a profile memory, not an evaluation report.",
        fill=LIGHT_BLUE,
    )


def add_project(doc: Document, project: dict) -> None:
    heading = doc.add_heading(project["name"], level=2)
    heading.paragraph_format.keep_with_next = True
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(5)
    meta.paragraph_format.keep_with_next = True
    set_font(meta.add_run(f"{project['category']} • github.com/tusharpathaknyu/{project['repository']}"), size=9.5, color=MUTED, italic=True)
    summary = doc.add_paragraph(project["summary"])
    summary.paragraph_format.keep_with_next = True
    for highlight in project["highlights"]:
        doc.add_paragraph(highlight, style="List Bullet")
    add_label_paragraph(doc, "Technologies", project["technology_context"])
    add_label_paragraph(doc, "Reusable resume wording", project["canonical_resume_bullet"], fill=LIGHT_BLUE)


def build_docx(context: dict) -> Path:
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_repeat_header_footer(doc)
    add_title_block(doc)

    doc.add_heading("Profile snapshot", level=1)
    for label, value in (
        ("Resume name", IDENTITY["resume_name"]),
        ("Email", IDENTITY["email"]),
        ("Phone", IDENTITY["phone"]),
        ("LinkedIn", IDENTITY["linkedin"]),
        ("GitHub", IDENTITY["github"]),
        ("Portfolio", IDENTITY["portfolio"]),
    ):
        add_label_paragraph(doc, label, value)

    doc.add_heading("Education", level=1)
    for item in EDUCATION:
        p = doc.add_paragraph()
        set_font(p.add_run(item["institution"]), bold=True, color=DARK_BLUE)
        set_font(p.add_run(f" — {item['degree']} | {item['dates']} | {item['location']}"))
        for detail in item["additional_context"]:
            doc.add_paragraph(detail, style="List Bullet")

    doc.add_heading("Work experience", level=1)
    for item in context["experience"]:
        doc.add_heading(item["organization"], level=2)
        role = add_label_paragraph(doc, "Role", f"{item['role']} | {item['dates']} | {item['location']}")
        role.paragraph_format.keep_with_next = True
        contributions = item["responsibilities_and_contributions"]
        for index, contribution in enumerate(contributions):
            bullet = doc.add_paragraph(contribution, style="List Bullet")
            bullet.paragraph_format.keep_with_next = index < len(contributions) - 1

    doc.add_page_break()
    doc.add_heading("Target role lanes represented in the current profile", level=1)
    for lane in TARGET_LANES:
        doc.add_paragraph(lane, style="List Bullet")

    doc.add_heading("Resume house style", level=1)
    add_label_paragraph(doc, "Format", "One-page, US Letter LaTeX resume using Charter at 9 pt with compact margins and section rules.")
    add_label_paragraph(doc, "Default order", "Education → Work Experience → role-specific Projects → Technical Skills.")
    for rule in context["resume_style"]["selection_rules"]:
        doc.add_paragraph(rule, style="List Bullet")

    doc.add_heading("Recent resume variants reviewed", level=2)
    for variant in context["resume_style"]["recent_resume_variants"]:
        p = doc.add_paragraph()
        set_font(p.add_run(f"{variant['timestamp']} — {variant['target_lane']}: "), bold=True, color=DARK_BLUE)
        set_font(p.add_run(f"experience order {variant['experience_order']}; projects {variant['selected_projects']}."))

    doc.add_page_break()
    doc.add_heading("Skill context", level=1)
    for group, skills in SKILL_GROUPS.items():
        add_label_paragraph(doc, group.replace("_", " ").title(), ", ".join(skills))

    doc.add_heading("Project portfolio", level=1)
    current_category = None
    for project in context["projects"]:
        if project["category"] != current_category:
            current_category = project["category"]
            doc.add_heading(current_category, level=1)
        add_project(doc, project)

    doc.add_heading("Profile details to add over time", level=1)
    doc.add_paragraph(
        "The context is designed to grow. These remaining preferences are useful for job filtering and application-form handling."
    )
    for item in context["profile_fields_to_collect_later"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Cross-check", level=1)
    doc.add_paragraph(
        "Review names, dates, responsibilities, project descriptions and preferred wording. Corrections can be applied once to the canonical JSON so every later resume variant uses the same profile memory."
    )
    doc.add_paragraph("Reviewer: ______________________________    Date: __________________")

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    return DOCX_PATH


def main() -> None:
    PROFILE_DIR.mkdir(parents=True, exist_ok=True)
    if SOURCE_TEMPLATE.is_file():
        TEMPLATE_PATH.write_text(SOURCE_TEMPLATE.read_text(encoding="utf-8"), encoding="utf-8")
    context = build_context()
    CONTEXT_PATH.write_text(json.dumps(context, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(CONTEXT_PATH)
    print(TEMPLATE_PATH)
    print(build_docx(context))


if __name__ == "__main__":
    main()
