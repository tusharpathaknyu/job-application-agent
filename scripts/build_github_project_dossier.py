from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts" / "Master_Candidate_Evidence_Dossier_Tushar_Pathak.docx"
METADATA = Path("/private/tmp/tushar_graphql.json")
ACTION_DIR = Path("/private/tmp/github_project_evidence/actions")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(35, 42, 48)
MUTED = RGBColor(92, 101, 110)
GREEN = RGBColor(33, 115, 70)
GOLD = RGBColor(145, 104, 20)
RED = RGBColor(155, 28, 28)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GOLD = "FFF6DB"
PALE_RED = "FDECEC"
PALE_GREEN = "EAF5EE"


PROJECTS = [
    {
        "name": "FPGA-Controlled Buck Converter",
        "repo": "fpga-buck-converter",
        "category": "Digital hardware, FPGA & validation",
        "posture": "Strong, scoped to completed bench phases",
        "signal": "mixed",
        "summary": "Digitally controlled 12 V-to-5 V buck-converter platform for a Gowin GW5A-25 FPGA, combining Verilog PWM, MCP3202 SPI acquisition and fixed-point PI control with staged bench bring-up.",
        "highlights": [
            "Repository reports five RTL testbenches with 32 checks and documents a signed-arithmetic bug found during closed-loop simulation.",
            "Bench log records FPGA/toolchain bring-up plus a scope-verified 100 kHz, 41.6% PWM output and a 0-to-5 V TC4420 driver output.",
            "The README explicitly distinguishes completed PWM/driver validation from unpowered MOSFET, open-loop buck, ADC and closed-loop hardware phases.",
        ],
        "tech": "Verilog, Icarus Verilog, Gowin EDA/CLI, Tcl, Make, SPI, fixed-point PI control, mixed-signal bench validation",
        "resume": "Designed and simulated an FPGA-based digital power controller with 100 kHz PWM, SPI ADC interface and fixed-point PI loop; validated FPGA and 5 V gate-driver stages on an oscilloscope through a staged hardware bring-up plan.",
        "verify": "Do not claim a working 12 V-to-5 V converter or closed-loop regulation yet; those hardware phases are documented as planned/RTL-only.",
        "sources": "README.md, Makefile, RTL/testbench tree, bench log; no GitHub Actions runs.",
    },
    {
        "name": "RV32I 5-Stage Pipelined Processor",
        "repo": "RISCV",
        "category": "Digital hardware, FPGA & validation",
        "posture": "Resolve CI failure before using strongest wording",
        "signal": "conflict",
        "summary": "Clean-room SystemVerilog implementation of a 32-bit, five-stage RV32I core with forwarding, load-use interlock, branch flushing, self-checking assembly tests, synthesis and gate-level simulation flows.",
        "highlights": [
            "README reports a 15-module synthesizable hierarchy, six self-checking programs passing at RTL and gate level, and a dependency-free Python assembler.",
            "Reported Yosys results are approximately 7,235 cells and 1,563 flip-flops; measured CPI ranges from 1.13 for straight-line logic to 1.65 for branch-heavy code.",
            "Latest GitHub Actions run on 2026-06-05 failed at the RTL self-checking regression, conflicting with the README's fully verified claim.",
        ],
        "tech": "SystemVerilog, Verilog, RISC-V RV32I, Python assembler, Icarus Verilog, Verilator, Yosys, Vivado Tcl, gate-level simulation",
        "resume": "Implemented a five-stage RV32I processor with data forwarding, load-use stalls and branch flushing, plus self-checking assembly, RTL synthesis and post-synthesis simulation flows.",
        "verify": "Re-run/fix the current CI regression and confirm whether contributor handle tpathak1411 is your prior account before claiming all implementation history.",
        "sources": "README.md, RTL/test trees, Makefile, commit/contributor history, latest CI run.",
    },
    {
        "name": "High-Speed SerDes Link Validation",
        "repo": "serdes-link-validation",
        "category": "Digital hardware, FPGA & validation",
        "posture": "Strong artifact-backed simulation project",
        "signal": "verified",
        "summary": "PyBERT/scikit-rf validation workflow for a measured backplane channel, including mixed-mode S-parameter analysis, CTLE/FFE/DFE equalization, eye/BER measurements and automated rate sweeps.",
        "highlights": [
            "Committed sweep_rate.csv records zero bit errors from 12 through 28 Gb/s and a 117.2 mV eye at 25.78 Gb/s.",
            "Workflow derives Rx FFE taps from measured post-CTLE pulse-response cursors and generates reproducible CSV/plot outputs.",
            "README explains that CTLE+DFE alone leaves the eye closed because pre-cursor ISI requires FFE correction.",
        ],
        "tech": "Python, PyBERT 9.2.1, scikit-rf, S-parameters, CTLE, Tx/Rx FFE, DFE, CDR, PRBS, BER/eye analysis",
        "resume": "Built an automated SerDes validation pipeline for a measured backplane channel, recovering error-free 12–28 Gb/s NRZ operation through channel-derived FFE plus CTLE/DFE equalization.",
        "verify": "README text says 122 mV at 25.78 Gb/s while the committed CSV says 117.2 mV; select one reproducible run and use its exact value.",
        "sources": "README.md, results/sweep_rate.csv, result figures, source/config tree; no CI workflow.",
    },
    {
        "name": "Wireless BMS Virtual HIL Tester",
        "repo": "wBMS-HIL-Tester",
        "category": "Digital hardware, FPGA & validation",
        "posture": "Strong, artifact- and CI-backed",
        "signal": "verified",
        "summary": "Pure-Python virtual hardware-in-the-loop framework that separates battery plant, burst-error channel and firmware state machine to validate wireless BMS framing, safety behavior and timing budgets.",
        "highlights": [
            "README reports 86 tests covering PEC15 vectors, frame round trips, ECM/PyBaMM plant physics, Hypothesis safety invariants and timing budgets; latest CI completed successfully.",
            "Committed channel report supports 2,816 packets, 165 corrupted packets, 517 flipped bits, 135 PEC rejections and a fault transition at 7.030 s.",
            "Test suite documented two design findings: an EMA coefficient that violated a 50 ms over-temperature budget near threshold and a watchdog/telemetry-cadence incompatibility.",
        ],
        "tech": "Python, pytest, Hypothesis, PyBaMM, battery ECM/SPMe models, Gilbert-Elliott channel, PEC15, property/timing testing",
        "resume": "Developed a virtual HIL framework for wireless BMS validation, exercising PEC15 framing, burst-error channels and safety timing; surfaced filter and watchdog assumptions through property-based and timing tests.",
        "verify": "Confirm whether ASIL-D wording should describe a modeled target/budget rather than formal safety certification.",
        "sources": "README.md, reports/channel_stats.txt, HTML/plot artifacts, tests, pyproject.toml, successful latest CI.",
    },
    {
        "name": "UVMForge",
        "repo": "UVMForge",
        "category": "Digital hardware, FPGA & validation",
        "posture": "Strong generator project; benchmark time claims",
        "signal": "verified",
        "summary": "Python/Jinja system that converts natural-language or imported hardware specifications into protocol-specific UVM environments, with RTL-aware port extraction, coverage-gap analysis and SVA generation.",
        "highlights": [
            "Repository includes templates and examples for APB, AXI4-Lite, UART, SPI and I2C, plus parsers for RTL, IP-XACT, SystemRDL, CSV and JSON inputs.",
            "Source tree contains dedicated coverage analysis, SVA generation and RTL-aware tests; latest GitHub Actions test workflow succeeded.",
            "Project includes a Streamlit UI, Docker/Cloud Run deployment files and multiple LLM provider paths including local Ollama.",
        ],
        "tech": "Python, Jinja2, UVM/SystemVerilog, SVA, RTL parsing, IP-XACT, SystemRDL, Streamlit, Docker, Cloud Run, Gemini/OpenAI/Ollama",
        "resume": "Built an RTL-aware UVM testbench generator supporting five bus protocols, specification import, coverage-gap analysis and SVA generation through reusable Python/Jinja templates.",
        "verify": "Time-saving and 'production-ready' claims are not backed by a committed benchmark; quantify generated-code compile rate and human edit time before using them.",
        "sources": "README.md, templates/examples/src/tests tree, workflows, successful latest test run.",
    },
    {
        "name": "RL-Based RTL Verification",
        "repo": "rl-rtl-verification",
        "category": "Digital hardware, FPGA & validation",
        "posture": "Promising; reproduce performance metrics",
        "signal": "mixed",
        "summary": "Reinforcement-learning bug hunter that wraps FIFO RTL as Gymnasium environments and uses PPO/curiosity-driven exploration, with an optional Verilator-backed simulation path.",
        "highlights": [
            "Nested project README reports a 100% bug-find rate versus a 0.5% random baseline and an optimal nine-step overflow-triggering sequence.",
            "Repository contains Python and real RTL/Verilator environments, multi-objective rewards, curiosity modules and training/evaluation utilities.",
            "Several advanced methods are present as modules, but the repository has no GitHub Actions history or consolidated reproducibility artifact.",
        ],
        "tech": "Python, PyTorch, Gymnasium, Stable-Baselines3 PPO, RND curiosity, Verilator, Verilog, TensorBoard",
        "resume": "Trained a PPO-based RTL bug-finding agent that learned a minimal nine-operation FIFO overflow sequence and integrated the environment with compiled Verilator simulation.",
        "verify": "Re-run the baseline/agent evaluation with fixed seeds and distinguish implemented experiments from advanced modules that are present but not benchmarked.",
        "sources": "project1_bug_sequence/README.md, RTL/RL/config tree, requirements; no CI runs.",
    },
    {
        "name": "WaveformGPT",
        "repo": "WaveformGPT",
        "category": "Digital hardware, FPGA & validation",
        "posture": "Feature-rich; latest CI failed",
        "signal": "conflict",
        "summary": "Python toolkit for querying VCD/FST simulation waveforms with pattern rules or LLMs, plus temporal assertions, protocol checks, regression comparison, visualization and live/voice interfaces.",
        "highlights": [
            "Repository implements VCD/FST parsing, AXI4/Wishbone/SPI checking, SVA-like temporal assertions, waveform comparison and CSV/JSON/WaveDrom/SystemVerilog/cocotb export.",
            "Demo/source tree includes live WebSocket/simulation monitoring and optional voice/LLM providers.",
            "Latest GitHub Actions workflow run on 2026-01-23 failed and exposed no job details through the API.",
        ],
        "tech": "Python, VCD/FST, temporal assertions, protocol checking, WebSockets, voice I/O, OpenAI/Gemini/Claude/Ollama, CI integration",
        "resume": "Created a natural-language waveform analysis toolkit with streaming parsers, protocol checks, temporal assertions and regression comparison across simulation traces.",
        "verify": "Fix or explain the latest CI failure and benchmark large-file parsing before claiming production-scale streaming.",
        "sources": "README.md, pyproject.toml, src/tests/demo tree, latest failed workflow run.",
    },
    {
        "name": "UART UVM Verification",
        "repo": "uart-verification",
        "category": "Digital hardware, FPGA & validation",
        "posture": "Useful portfolio project; coverage report needs raw evidence",
        "signal": "mixed",
        "summary": "SystemVerilog/UVM environment for UART TX/RX verification with sequences, driver, monitor, scoreboard, coverage model and deliberately buggy RTL variants.",
        "highlights": [
            "Coverage report claims 95.2% overall functional coverage, 97.1% line, 94.6% branch, 100% FSM and 91.0% toggle coverage.",
            "Verification plan targets four injected defects: baud counter, parity calculation, stop-bit checking and sample-point errors.",
            "Repository provides UVM and EDA Playground variants but no CI run or raw simulator coverage database.",
        ],
        "tech": "SystemVerilog, UVM 1.2, functional/code/FSM/toggle coverage, assertions, EDA Playground, Make",
        "resume": "Built a UVM environment for UART verification with constrained sequences, scoreboarding, error injection and a documented 95.2% functional-coverage closure plan.",
        "verify": "Regenerate the coverage report from a named simulator and retain raw logs before presenting coverage percentages as independently reproduced results.",
        "sources": "README.md, docs/coverage_report.md, RTL/testbench/EDA Playground tree; no CI runs.",
    },
    {
        "name": "ARCS — Autoregressive Circuit Synthesis",
        "repo": "ARCS",
        "category": "AI/ML for circuits & engineering",
        "posture": "Strong research project with result artifacts",
        "signal": "verified",
        "summary": "Spec-conditioned circuit-generation research code using native component/value tokens, autoregressive transformers, constrained decoding, graph generative models and SPICE-based ranking/RL.",
        "highlights": [
            "Committed publication_eval_v3.json supports a 6.432 mean hybrid reward, 99.875% simulation validity and approximately 0.05 s wall time across the reported setup.",
            "README reports GRPO improving the graph-transformer path to 96.6% structural and 53.1% simulation validity over five seeds.",
            "Repository contains extensive ablations, multi-seed outputs, paper assets and a latest successful consistency workflow.",
        ],
        "tech": "Python, PyTorch, autoregressive transformers, graph models, VAE, conditional flow matching, GRPO/REINFORCE, ngspice, LaTeX",
        "resume": "Developed a spec-conditioned circuit generator using component-level tokens and SPICE-grounded ranking, achieving 99.875% simulation validity in the committed hybrid evaluation artifact.",
        "verify": "Define the evaluation reward in interview/resume supporting material; avoid 'first' claims unless backed by a completed literature review/publication.",
        "sources": "README.md, ARCHITECTURE.md, results/publication_eval_v3.json and other result artifacts, successful latest CI.",
    },
    {
        "name": "Neural Surrogate + RL Circuit Design",
        "repo": "neural-surrogate-rl-circuits",
        "category": "AI/ML for circuits & engineering",
        "posture": "Resolve conflicting result narratives",
        "signal": "conflict",
        "summary": "Power-converter design system combining a learned waveform surrogate, PPO policies and periodic ngspice validation across seven converter topologies.",
        "highlights": [
            "Current README reports all seven topologies below SPICE MSE 5 using a dual surrogate/formula architecture and 268K+ SPICE simulations.",
            "Committed TRAINING_RESULTS.md from 2026-02-12 instead states only 4/7 achieved MSE below 5 and flags buck-boost/flyback generalization failures.",
            "Repository contains trained checkpoints for all seven topologies, training notebooks/scripts and web/Hugging Face deployment assets.",
        ],
        "tech": "Python, PyTorch, PPO actor-critic, neural waveform surrogate, ngspice, analytical converter models, Jupyter, Hugging Face",
        "resume": "Built a SPICE-in-the-loop RL pipeline for multi-topology power-converter tuning using a fast neural surrogate and per-topology policies.",
        "verify": "Reconcile README and TRAINING_RESULTS metrics, identify the exact checkpoint/evaluation script for each value, and only then add all-seven-topology or 100,000× speed claims.",
        "sources": "README.md, TRAINING_RESULTS.md, checkpoint/training tree; no CI runs.",
    },
    {
        "name": "PowerElecLLM",
        "repo": "PowerElecLLM",
        "category": "AI/ML for circuits & engineering",
        "posture": "Strong benchmark concept; reconcile dataset counts",
        "signal": "conflict",
        "summary": "Benchmark and SPICE-grounded evaluation framework for measuring LLM performance on power-electronics circuit design and preparing fine-tuning datasets.",
        "highlights": [
            "README reports four-model evaluation and 25.0% best accuracy for a fine-tuned GPT-4o, using output voltage within 10% of target as a pass criterion.",
            "Committed combined benchmark summary contains 500 problems (400 train/100 test), while the README headline says 650 and elsewhere also calls the benchmark 500 problems.",
            "Repository includes expert-verified GATE/MIT problems, synthetic SPICE cases, evaluation outputs, charts and fine-tuning formats; latest Python CI succeeded.",
        ],
        "tech": "Python, ngspice, LLM evaluation/fine-tuning, benchmark design, OpenAI/Gemini/LLaMA/Grok adapters, JSONL, statistical reporting",
        "resume": "Created a SPICE-grounded power-electronics benchmark with expert and synthetic problems, held-out evaluation splits and multi-provider LLM/fine-tuning pipelines.",
        "verify": "Reconcile 500 versus 650 problems and trace the 25.0%/2.3% accuracy numbers to exact committed result files and evaluation versions.",
        "sources": "README.md, benchmark_summary.json, benchmark/evaluation/chart tree, successful latest CI.",
    },
    {
        "name": "OhmAI",
        "repo": "OhmAI",
        "category": "AI/ML for circuits & engineering",
        "posture": "Resolve test artifact conflict",
        "signal": "conflict",
        "summary": "FastAPI application that converts natural-language requests into SPICE netlists using Gemini, validates circuits through ngspice and emits Datadog metrics/traces/alerts.",
        "highlights": [
            "Architecture combines FastAPI, Gemini 2.5, template/rule logic, ngspice measurements and Datadog observability with Docker/Cloud Run deployment files.",
            "README claims a 39/39 passing validation suite and reports per-circuit error distributions.",
            "Committed test_results.json contains 33 HTTP_ERROR entries (all HTTP 422), directly conflicting with the README test claim; no GitHub Actions runs are present.",
        ],
        "tech": "Python, FastAPI, Gemini/Vertex AI, ngspice, Datadog APM/StatsD/events, Docker, Cloud Run, Chart.js",
        "resume": "Developed a natural-language-to-SPICE service with simulator-backed validation and end-to-end Datadog observability for circuit-generation requests.",
        "verify": "Re-run both test suites, explain whether test_results.json is an obsolete negative test, and commit a dated passing report before using accuracy/pass-rate numbers.",
        "sources": "README.md, test_results.json, test scripts, backend/Datadog/deployment tree; no CI runs.",
    },
    {
        "name": "PCB Thermal AI Predictor",
        "repo": "PCB-Thermal-AI",
        "category": "AI/ML for circuits & engineering",
        "posture": "Good prototype; validate model metrics",
        "signal": "mixed",
        "summary": "U-Net-based predictor for 2D PCB temperature fields from copper, via, component and power maps, with ONNX export, uncertainty estimates, web demo and REST API paths.",
        "highlights": [
            "README reports 6.0 °C mean absolute error, approximately 17 °C maximum error and 27 ms ONNX inference for a 4.3M-parameter U-Net.",
            "Committed dataset_stats.json corroborates 2,000 synthetic samples with 1,600/200/200 train/validation/test splits and four input channels.",
            "Repository includes PyTorch training, augmentation, MC-dropout uncertainty, FastAPI/Streamlit deployment and thermal visualization utilities.",
        ],
        "tech": "Python, PyTorch U-Net, ONNX, NumPy/Pandas/OpenCV, FastAPI, Streamlit, synthetic thermal simulation, MC Dropout",
        "resume": "Trained and deployed a U-Net thermal-field predictor on 2,000 synthetic PCB samples with ONNX inference and uncertainty-aware visualization.",
        "verify": "Commit an evaluation artifact that reproduces the 6.0 °C MAE, ~17 °C max error and 27 ms latency; dataset statistics alone do not validate model accuracy.",
        "sources": "README.md, dataset_stats.json, training/API/demo tree; no CI runs.",
    },
    {
        "name": "GymEZ",
        "repo": "GymEZ",
        "category": "Product, embedded & full-stack systems",
        "posture": "Major private prototype; latest CI failed",
        "signal": "conflict",
        "summary": "Private fitness/game product connecting camera-based pose controls and workout progression to a Unity combat game, backed by a React Native app and Supabase services; repository also contains CombatBand wearable design material.",
        "highlights": [
            "README states an end-to-end Android prototype works on a Samsung Galaxy S24, with per-hand punches, blocking, workout-driven stats and Gemini meal scanning.",
            "Repository documents React Native/TypeScript, Unity/C#, MediaPipe pose detection, Supabase, extensive services/screens and an nRF52840/BMI270/MAX30102 wearable concept/PCB.",
            "Latest CI on 2026-06-30 failed in Lint & Type Check; downstream unit/build/E2E jobs were skipped.",
        ],
        "tech": "React Native, TypeScript, Unity 6/C#, MediaPipe, Supabase/PostgreSQL/Realtime, Gemini, BLE, KiCad, nRF52840, BMI270, MAX30102",
        "resume": "Built an Android fitness-game prototype that maps live camera pose events and workout history into a Unity combat system through a React Native/Supabase architecture.",
        "verify": "Fix lint/type errors; separate implemented wearable hardware from design documentation; confirm the 65-screen/44-service counts and any production-user claims before use.",
        "sources": "Private README.md, GYMEZ_EVERYTHING.md, IMPLEMENTATION_SUMMARY.md, package manifest, code/hardware/game tree, latest failed CI.",
    },
    {
        "name": "GuideBot",
        "repo": "guidebot",
        "category": "Product, embedded & full-stack systems",
        "posture": "Private full-stack prototype; add tests",
        "signal": "mixed",
        "summary": "Voice-and-vision AI workbench assistant for hardware makers with live camera/whiteboard context, structured schematic drawing and pay-as-you-go SaaS billing.",
        "highlights": [
            "Code implements Flask-SocketIO sessions, OpenAI text/vision/speech/image calls and a structured draw_schematic tool for circuit diagrams.",
            "SaaS layer includes Google OAuth, Stripe top-ups/webhooks, a prepaid wallet, usage ledger and SQLite/PostgreSQL support.",
            "Repository contains Render deployment configuration but no README, tests or GitHub Actions history.",
        ],
        "tech": "Python, Flask, Flask-SocketIO, OpenAI multimodal APIs, JavaScript/HTML/CSS, SQLAlchemy, Google OAuth, Stripe, PostgreSQL/SQLite, Render",
        "resume": "Developed a multimodal hardware-workbench assistant with live voice/vision context, structured schematic generation and transparent usage-based billing.",
        "verify": "Add a README, tests and threat/privacy review; verify the configured markup percentage because source comments and constants disagree (3% versus 5%).",
        "sources": "Private server.py, models.py, requirements.txt, DEPLOY.md, frontend tree; no CI runs.",
    },
    {
        "name": "GymEZ Fighting Game Prototype",
        "repo": "gymez-fighting-game",
        "category": "Product, embedded & full-stack systems",
        "posture": "Historical prototype; merge into GymEZ story",
        "signal": "mixed",
        "summary": "Early Godot 4.5/HTML5 prototype that validated camera-as-controller combat before the production direction moved to Unity in the main GymEZ repository.",
        "highlights": [
            "README says camera_pose.js and the camera-controller concept carried into the current product while the Godot engine and CombatBand direction changed.",
            "Repository includes GDScript game logic, JavaScript camera pose detection and a React Native bridge prototype.",
            "README labels the project archived, but GitHub's repository archived flag is currently false.",
        ],
        "tech": "Godot 4.5, GDScript, JavaScript pose detection, HTML5/WebView, Python bridge code",
        "resume": "Prototyped camera-driven combat in Godot/HTML5 and used the result to de-risk pose controls before migrating the production game to Unity.",
        "verify": "Either archive the repository in GitHub or remove the mismatch; present it as a validation milestone within GymEZ, not as a separate active product.",
        "sources": "README.md and project/source tree; no CI runs.",
    },
    {
        "name": "Simulink Magnetic Signal Analyzer",
        "repo": "simulink-magnetic-analyzer",
        "category": "Product, embedded & full-stack systems",
        "posture": "Small hardware demo; verify live results",
        "signal": "mixed",
        "summary": "ESP32/SS49E-to-Simulink signal-processing demo for live magnetic-field visualization, filtering, statistics and threshold alerts.",
        "highlights": [
            "README documents a 500 Hz sensor stream over 115200-baud serial into MATLAB/Simulink.",
            "Planned model computes moving-average filtering, mean/RMS/peak-to-peak metrics and dashboard alerts.",
            "Repository has a single commit and includes a MATLAB serial test plus a linked external demo video, but no committed Simulink model artifact is visible in the tree.",
        ],
        "tech": "MATLAB, Simulink, ESP32, SS49E Hall sensor, serial acquisition, moving-average filtering, real-time dashboards",
        "resume": "Integrated an ESP32 Hall-sensor stream with MATLAB/Simulink for live filtering, statistics and magnetic-signal visualization.",
        "verify": "Confirm the linked demo and retain the actual .slx model/results in the repository before describing the complete model as reproducible.",
        "sources": "README.md, matlab/testSerialConnection.m, repository tree; no CI runs.",
    },
    {
        "name": "Magnetic SaaS / MFL Fingerprinting",
        "repo": "magnetic-saas-mfl",
        "category": "Product, embedded & full-stack systems",
        "posture": "Concept-to-pipeline prototype; results not recorded",
        "signal": "mixed",
        "summary": "Non-contact magnetic load-state/anomaly pipeline using an ESP32 Hall sensor, Python collection/feature extraction, classical ML and a Streamlit dashboard.",
        "highlights": [
            "Repository includes firmware, live collector, feature extraction, model training and dashboard paths.",
            "README defines idle/medium/high load collection workflow and safe, non-invasive sensor placement.",
            "The stated 80–90% classification accuracy is a success criterion, not a reported achieved metric.",
        ],
        "tech": "ESP32/Arduino C++, Hall sensing, Python, NumPy/Pandas/SciPy, scikit-learn, FFT/features, anomaly detection, Streamlit",
        "resume": "Built an end-to-end non-contact magnetic-fingerprinting prototype spanning ESP32 acquisition, feature extraction, load-state classification and live monitoring.",
        "verify": "Collect and commit labeled data, confusion matrices and anomaly results before citing any accuracy or false-positive number.",
        "sources": "README.md, firmware/collector/analysis/dashboard tree, requirements; no CI runs.",
    },
    {
        "name": "IPC, Threading & Scheduling",
        "repo": "ipc_threading_scheduling",
        "category": "Systems & foundational projects",
        "posture": "Older coursework; verify authorship and execution",
        "signal": "mixed",
        "summary": "C program that partitions a large numeric input among POSIX threads, exchanges data through System V shared memory and coordinates read/sum phases with mutexes/condition variables and scheduling logic.",
        "highlights": [
            "Single-file implementation uses pthreads, shared memory, process control and timing instrumentation.",
            "Repository has one commit attributed to tpathak1411 rather than the currently authenticated GitHub handle.",
        ],
        "tech": "C, POSIX threads, System V shared memory, mutexes/condition variables, process scheduling, file partitioning",
        "resume": "Parallelized large-file ingestion and aggregation in C using pthreads and shared-memory IPC, with synchronization and timing instrumentation.",
        "verify": "Confirm tpathak1411 is your previous account, compile with warnings enabled and document input size/runtime comparisons before using performance language.",
        "sources": "priority-scheduling-main.c, repository/contributor metadata; no README or CI.",
    },
    {
        "name": "Multithreaded NxN Sudoku Solver",
        "repo": "Sudoku",
        "category": "Systems & foundational projects",
        "posture": "Older coursework; verify build and correctness",
        "signal": "mixed",
        "summary": "C backtracking solver that creates pthread branches for candidate values up to a thread cap and supports square Sudoku grids up to 36×36.",
        "highlights": [
            "Source contains recursive safety checks, branch-local grid copies and a shared solution flag.",
            "Repository has one commit attributed to tpathak1411, no README, tests or CI, and the source references MAX_THREADS without an active definition in the committed file.",
        ],
        "tech": "C, pthreads, recursive backtracking, concurrency, file input",
        "resume": "Explored parallel backtracking for generalized NxN Sudoku by spawning bounded pthread branches across valid candidate assignments.",
        "verify": "Confirm old-handle authorship, fix the MAX_THREADS build issue and add correctness/performance tests before listing as a completed solver.",
        "sources": "sudoku.c and repository/contributor metadata; no README or CI.",
    },
]


NON_PROJECTS = [
    {
        "name": "tusharpathaknyu.github.io",
        "kind": "Private portfolio website",
        "note": "Static HTML/CSS/JavaScript portfolio that presents education, experience, project cards, tech stack and three targeted resume downloads. Latest Pages deployment workflow succeeded on 2026-06-11. Treat biography/employment claims as candidate-provided, not repository-verified evidence.",
    },
    {
        "name": "tusharpathaknyu",
        "kind": "Public GitHub profile README",
        "note": "Profile summary listing NYU/BITS affiliation, hardware/AI positioning, a compact stack and selected projects. Update stale project naming/status before using it as a canonical source.",
    },
    {
        "name": "react-native-health-connect",
        "kind": "Public fork",
        "note": "Fork of a React Native Android Health Connect library. GitHub ownership alone does not establish contribution; no unique commits or changes were assessed in this review.",
    },
    {
        "name": "awesome-opensource-hardware",
        "kind": "Public fork",
        "note": "Forked open-source hardware resource list. Keep out of the original-project section unless you can point to upstreamed contributions or substantive fork-specific work.",
    },
]


DISCREPANCIES = [
    ("GymEZ", "Latest CI failed in Lint & Type Check; all later jobs skipped.", "Fix lint/type errors and rerun full unit/build/E2E workflow."),
    ("RISCV", "README says all RTL/gate tests pass; latest CI failed RTL regression.", "Reproduce locally, fix CI, then cite exact passing commit."),
    ("WaveformGPT", "Latest workflow failed; API exposes no job detail.", "Repair workflow and add explicit parser/protocol regression results."),
    ("Neural Surrogate + RL", "README says 7/7 SPICE MSE <5; TRAINING_RESULTS says 4/7.", "Version results/checkpoints and publish one authoritative table."),
    ("PowerElecLLM", "README alternates between 650 and 500 problems; committed summary says 500.", "Reconcile corpus definition and link each accuracy to a result file."),
    ("OhmAI", "README claims 39/39 passing; committed test_results has 33 HTTP 422 errors.", "Identify stale/negative suite, rerun, and commit dated passing evidence."),
    ("SerDes", "README says 122 mV at 25.78 Gb/s; sweep CSV records 117.2 mV.", "Choose one config/run and use its exact artifact value."),
    ("GymEZ Godot prototype", "README says archived; repository archived flag is false.", "Archive the repo or update status wording."),
    ("UART verification", "Coverage percentages appear only in a Markdown report.", "Retain simulator logs/database and document tool/version/command."),
    ("Sudoku + IPC", "Only commit author is tpathak1411, not current handle.", "Confirm this is your previous account before claiming authorship."),
]


WORK_EXPERIENCE = [
    {
        "organization": "LayeredAI",
        "role": "AI Solutions Engineer",
        "dates": "March 2026 - Present",
        "location": "Remote",
        "evidence": "Resume + private portfolio only",
        "confidence": "Needs employer/engagement evidence",
        "safe": [
            "Customer-facing engineering role translating stakeholder requirements into technical recommendations and custom AI solutions.",
            "Creates technical documentation, demos and presentations for technical and non-technical audiences.",
        ],
        "verify": "Confirm legal employer/client relationship, exact start date, production deployments, tools/languages and measurable delivery outcomes.",
    },
    {
        "organization": "Dreamline AI",
        "role": "Software Engineer, Software Development Team",
        "dates": "August 29, 2025 - April 23, 2026",
        "location": "Remote",
        "evidence": "Employer experience certificate supplied by candidate",
        "confidence": "Certificate-backed",
        "safe": [
            "Designed and developed scalable software applications while following coding standards, best practices and project requirements.",
            "Implemented and maintained RESTful APIs and backend services and improved data-processing workflows for reliability and efficiency.",
            "Collaborated with product, UI/UX and QA teams and participated in code reviews focused on quality and maintainability.",
        ],
        "verify": "C++/Python, CI/CD, regression automation and modeling claims appear in resumes but not in the supplied certificate; retain only with code, manager confirmation or work samples.",
    },
    {
        "organization": "Texas Instruments (India) Pvt. Ltd.",
        "role": "Trainee / Applications Engineering Intern (resume wording)",
        "dates": "January 23, 2023 - July 14, 2023",
        "location": "Bangalore, India",
        "evidence": "Training certificate details supplied by candidate",
        "confidence": "Dates and project titles certificate-backed",
        "safe": [
            "Completed a training project on competitive benchmarking and teardown of a charger under project guide Ramkumar S.",
            "Completed PCB board design work in Altium as part of the training engagement.",
        ],
        "verify": "The certificate text provided does not establish 65 W GaN/flyback/USB-PD specifics, Python automation, test instruments, FLIR, four-layer layout, bring-up or measured results. Link those details to a report or manager confirmation before use.",
    },
    {
        "organization": "Xarvis",
        "role": "Software Developer Intern",
        "dates": "May 2021 - June 2021",
        "location": "Not supplied",
        "evidence": "Candidate statement + private portfolio",
        "confidence": "Needs certificate/reference",
        "safe": [
            "Extended domain and repository layers for a new tenant-information schema.",
            "Implemented add, delete and get REST endpoints in the controller layer and integrated them with Swagger for discoverable, executable API documentation.",
        ],
        "verify": "Confirm employer name styling, exact dates, technology stack, team context and whether endpoints were deployed to production.",
    },
]


RESUME_VARIANTS = [
    ("2026-06-29 19:13", "Power-conversion architecture", "TI, LayeredAI, Dreamline", "FPGA buck; neural surrogate; GymEZ CombatBand", "Latest LaTeX pattern; visually clean but contains multiple unsafe hardware-result claims."),
    ("2026-06-29 19:00", "Analog power / board hardware", "TI, LayeredAI, Dreamline", "FPGA buck; GymEZ CombatBand; neural surrogate", "Strong analog-applications framing; same unsupported buck/board bring-up claims."),
    ("2026-06-28 20:40", "SoC/ASIC verification", "LayeredAI, Dreamline, TI", "RISCV; UVMForge; ARCS", "Good project selection; overstates RISCV CI/timing closure and some tool proficiency."),
    ("2026-06-28 19:21", "SoC/ASIC verification", "LayeredAI, Dreamline, TI", "RISCV; UVMForge; ARCS", "Near-duplicate of prior version with condensed TI section and extra imaging/CV-adjacent wording."),
    ("2026-06-27 17:05", "SerDes / mixed-signal validation", "TI, LayeredAI, Dreamline", "SerDes; FPGA buck; GymEZ CombatBand", "SerDes evidence is strongest; buck and CombatBand bench claims need correction."),
]


RESUME_RISKS = [
    ("FPGA buck converter", "Claims a complete 12 V-to-5 V converter, closed-loop regulation, ripple and load-step validation.", "GitHub bench log currently supports FPGA PWM and a 5 V TC4420 driver only; later power stages are planned/RTL-only.", "Block from tailoring until hardware evidence is added."),
    ("Neural surrogate + RL", "Claims 268K+ simulations, 1,000x+ speedup and successful multi-topology optimization.", "Repository result narratives conflict on 7/7 versus 4/7 topologies; speedup is not tied to a reviewed benchmark artifact.", "Use architecture/process wording; omit performance numbers pending reconciliation."),
    ("GymEZ CombatBand", "Claims a completed four-layer board, full bring-up, battery subsystem and signal-integrity checks.", "Private repository contains extensive design/PCB documentation, but reviewed evidence did not establish completed physical bring-up.", "Describe as PCB/system design unless bench artifacts are provided."),
    ("RISCV", "Claims full passing CI, post-synthesis equivalence, Artix-7 timing closure and power-optimized implementation.", "Latest CI fails at RTL regression; README supports synthesis/GLS methodology but not every optimization claim.", "Fix CI and retain exact tool reports before using strongest language."),
    ("UVMForge", "Claims an 80-test regression, deployment and 50+ SVA patterns.", "Latest test workflow passed and source supports the feature set, but a dated raw test-count/compile benchmark was not reviewed.", "Safe at feature level; qualify exact counts until a report is committed."),
    ("Texas Instruments", "Adds 65 W GaN flyback USB-PD, detailed instruments, Python automation, FLIR and four-layer board bring-up.", "Supplied certificate supports charger benchmarking/teardown and Altium PCB design, not those implementation details.", "Keep certificate-backed wording unless project report or manager confirmation supports expansion."),
    ("Dreamline AI", "Adds C++/Python, CI/CD, regression automation and design-flow/modeling work.", "Certificate supports scalable applications, REST/backend work, data workflows, cross-functional delivery and code reviews.", "Use certificate wording by default; add technologies only with work evidence."),
    ("LayeredAI", "Claims production tools/deployments and concurrent customer engagements.", "Only candidate-authored resume/portfolio evidence was reviewed.", "Obtain employer/client-safe evidence and measurable outcomes."),
    ("Technical skills", "Lists commercial EDA tools, HVDC/server-power concepts, SIMPLIS and other domain concepts alongside practiced skills.", "Repositories and work evidence do not establish equal hands-on depth for every item.", "Split into hands-on, working knowledge and conceptual familiarity; tailor conservatively."),
    ("Recognition heading", "Uses 'Patents & Recognition' while listing no patent.", "No patent evidence was supplied.", "Rename to 'Recognition' unless an actual patent/application is documented."),
]


def load_metadata() -> dict[str, dict]:
    raw = json.loads(METADATA.read_text(encoding="utf-8"))
    nodes = raw["data"]["user"]["repositories"]["nodes"]
    return {node["name"]: node for node in nodes}


def latest_action(repo: str) -> str:
    path = ACTION_DIR / f"{repo}.json"
    if not path.exists():
        return "No Actions run"
    runs = json.loads(path.read_text(encoding="utf-8")).get("workflow_runs", [])
    if not runs:
        return "No Actions run"
    run = runs[0]
    return f"{run.get('conclusion') or run.get('status')} ({run['updated_at'][:10]})"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str, color=BLUE) -> None:
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), str(color))
    r_pr.append(c)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph) -> None:
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Number"):
        style = doc.styles[list_style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_label_paragraph(doc, label: str, text: str, *, fill: str | None = None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    if fill:
        p_pr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        p_pr.append(shd)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "120")
        ind.set(qn("w:right"), "120")
        p_pr.append(ind)
    lead = paragraph.add_run(f"{label}: ")
    set_font(lead, bold=True, color=DARK_BLUE)
    body = paragraph.add_run(text)
    set_font(body)
    return paragraph


def format_table_text(table, header=True, size=9) -> None:
    for ridx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_font(run, size=size, bold=(header and ridx == 0))
            if header and ridx == 0:
                set_cell_shading(cell, LIGHT_BLUE)


def language_summary(meta: dict) -> str:
    edges = meta.get("languages", {}).get("edges", [])
    if not edges:
        return "Not detected"
    total = sum(item["size"] for item in edges) or 1
    return ", ".join(f"{item['node']['name']} {item['size'] / total:.0%}" for item in edges[:3])


def add_project(doc: Document, project: dict, meta: dict) -> None:
    heading = doc.add_paragraph(style="Heading 2")
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(project["name"])
    set_font(run, size=13, bold=True, color=BLUE)

    status_color = GREEN if project["signal"] == "verified" else RED if project["signal"] == "conflict" else GOLD
    status = doc.add_paragraph()
    status.paragraph_format.space_after = Pt(5)
    marker = status.add_run(project["posture"])
    set_font(marker, size=9.5, bold=True, color=status_color)
    status.add_run("  •  ")
    visibility = "Private" if meta.get("isPrivate") else "Public"
    commit_count = (((meta.get("defaultBranchRef") or {}).get("target") or {}).get("history") or {}).get("totalCount", "—")
    status_text = (
        f"{visibility} • {commit_count} default-branch commits • "
        f"created {meta.get('createdAt', '')[:10]} • last push {meta.get('pushedAt', '')[:10]} • "
        f"latest CI: {latest_action(project['repo'])}"
    )
    muted = status.add_run(status_text)
    set_font(muted, size=9.5, color=MUTED)

    link_line = doc.add_paragraph()
    link_line.paragraph_format.space_after = Pt(6)
    add_hyperlink(link_line, f"github.com/tusharpathaknyu/{project['repo']}", f"https://github.com/tusharpathaknyu/{project['repo']}")

    doc.add_paragraph(project["summary"])
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(3)
    label.paragraph_format.space_after = Pt(2)
    set_font(label.add_run("Evidence-backed highlights"), bold=True, color=DARK_BLUE)
    for item in project["highlights"]:
        doc.add_paragraph(item, style="List Bullet")

    add_label_paragraph(doc, "Technology evidence", project["tech"])
    add_label_paragraph(doc, "GitHub Linguist mix (may include generated/report assets)", language_summary(meta))
    fill = PALE_GREEN if project["signal"] == "verified" else PALE_RED if project["signal"] == "conflict" else PALE_GOLD
    add_label_paragraph(doc, "Resume-ready draft", project["resume"], fill=fill)
    add_label_paragraph(doc, "Cross-check before use", project["verify"])
    source = doc.add_paragraph()
    source.paragraph_format.space_before = Pt(2)
    source.paragraph_format.space_after = Pt(8)
    set_font(source.add_run("Evidence reviewed: "), size=9, bold=True, color=MUTED)
    set_font(source.add_run(project["sources"]), size=9, italic=True, color=MUTED)


def build() -> Path:
    metadata = load_metadata()
    doc = Document()
    setup_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("MASTER CANDIDATE EVIDENCE DOSSIER"), size=8.5, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Cross-verification draft  •  "), size=8.5, color=MUTED)
    add_page_field(footer)
    for run in footer.runs:
        set_font(run, size=8.5, color=MUTED)

    # Editorial-cover pattern, compact-reference override: reduced whitespace so evidence starts on page 1.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(30)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    set_font(kicker.add_run("RESUME & PORTFOLIO EVIDENCE REVIEW"), size=10, bold=True, color=GOLD)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(7)
    set_font(title.add_run("Master Candidate Evidence Dossier"), size=27, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)
    set_font(subtitle.add_run("Tushar Dhananjay Pathak | Work history, resume patterns & GitHub projects"), size=13.5, color=DARK_BLUE)
    meta_line = doc.add_paragraph()
    meta_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_line.paragraph_format.space_after = Pt(22)
    set_font(meta_line.add_run(f"Cross-verification draft • GitHub snapshot: {date(2026, 7, 1).strftime('%B %-d, %Y')}"), size=10, color=MUTED)

    add_label_paragraph(
        doc,
        "Purpose",
        "Consolidate work-history evidence, recent resume patterns and repository-backed project information for truthful resume tailoring. Draft bullets are intentionally conservative and should not be used until every highlighted cross-check is resolved.",
        fill=LIGHT_BLUE,
    )

    doc.add_heading("Executive review", level=1)
    for item in (
        "Source set reviewed: five newest Applications PDFs from Downloads, the supplied LaTeX template, candidate-provided Dreamline/Texas Instruments/Xarvis information, and the authenticated tusharpathaknyu GitHub account.",
        "All five resumes are one-page US Letter variants using the same dense Charter/section-rule system; they target analog power, SoC/ASIC verification, or SerDes/mixed-signal roles.",
        "Inventory: 24 repositories total — 22 original repositories and 2 forks. The dossier treats 20 as technical/product projects and separates the portfolio/profile repositories and forks.",
        "Most original repositories show a single contributor in GitHub metadata. RISCV also includes tpathak1411; Sudoku and IPC history is entirely under that handle and needs identity confirmation.",
        "Strongest evidence packages combine result artifacts and successful CI: ARCS, wBMS-HIL-Tester and UVMForge. SerDes has strong committed CSV/plot evidence but no CI.",
        "The most urgent resume risks are unsupported work details, incomplete hardware represented as validated, conflicting project metrics and failed latest CI runs. These are listed next and repeated in the relevant sections.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()
    doc.add_heading("Canonical candidate baseline", level=1)
    add_label_paragraph(doc, "Resume name", "Tushar Pathak")
    add_label_paragraph(doc, "Legal/full name seen in supplied records", "Pathak Tushar Dhananjay / Tushar Dhananjay Pathak")
    add_label_paragraph(doc, "Phone", "+1 (718) 413-9793")
    add_label_paragraph(doc, "Email", "tpathak1411@gmail.com")
    add_label_paragraph(doc, "LinkedIn", "linkedin.com/in/tushar-pathak-7b945219a")
    add_label_paragraph(doc, "GitHub", "github.com/tusharpathaknyu")
    add_label_paragraph(doc, "Portfolio", "tusharpathaknyu.github.io")
    add_label_paragraph(
        doc,
        "Contact verification",
        "These values are consistent across the reviewed resumes/template but were not independently verified against identity records. Use the short LinkedIn display text while retaining the full destination URL.",
        fill=PALE_GOLD,
    )

    doc.add_heading("Education baseline", level=2)
    for item in (
        "New York University — M.S. Computer Engineering, January 2024 - December 2025; resumes also state Graduate Scholarship Recipient.",
        "Birla Institute of Technology and Science (BITS) Pilani, Hyderabad — B.E. Electrical and Electronics Engineering, August 2019 - July 2023.",
    ):
        doc.add_paragraph(item, style="List Bullet")
    add_label_paragraph(doc, "Evidence status", "Repeated consistently across all five resumes; diploma/transcript and scholarship documentation were not supplied in this review.")

    doc.add_heading("Work-experience evidence matrix", level=1)
    doc.add_paragraph(
        "Use exact certificate dates where available. Month-only resume dates are acceptable for display, but the agent should retain the exact dates in its evidence store."
    )
    work_table = doc.add_table(rows=1, cols=4)
    work_table.style = "Table Grid"
    for i, text in enumerate(("Organization", "Role and dates", "Evidence reviewed", "Tailoring status")):
        work_table.rows[0].cells[i].text = text
    repeat_table_header(work_table.rows[0])
    for item in WORK_EXPERIENCE:
        row = work_table.add_row().cells
        row[0].text = item["organization"]
        row[1].text = f"{item['role']}\n{item['dates']}"
        row[2].text = item["evidence"]
        row[3].text = item["confidence"]
    set_table_geometry(work_table, [1700, 2650, 2500, 2510])
    for row in work_table.rows:
        prevent_row_split(row)
    format_table_text(work_table, size=8.8)

    for item in WORK_EXPERIENCE:
        doc.add_heading(item["organization"], level=2)
        detail = doc.add_paragraph()
        detail.paragraph_format.space_after = Pt(5)
        set_font(detail.add_run(f"{item['role']}  •  {item['dates']}  •  {item['location']}"), bold=True, color=DARK_BLUE)
        evidence = doc.add_paragraph()
        evidence.paragraph_format.space_after = Pt(4)
        set_font(evidence.add_run(f"Evidence: {item['evidence']}  |  Status: {item['confidence']}"), size=9.5, color=MUTED, italic=True)
        label = doc.add_paragraph()
        label.paragraph_format.space_after = Pt(2)
        label.paragraph_format.keep_with_next = True
        set_font(label.add_run("Conservative resume wording"), bold=True, color=DARK_BLUE)
        for bullet in item["safe"]:
            doc.add_paragraph(bullet, style="List Bullet")
        add_label_paragraph(doc, "Cross-check before expansion", item["verify"], fill=PALE_GOLD)

    doc.add_page_break()
    doc.add_heading("Recent resume pattern review", level=1)
    doc.add_paragraph(
        "The five newest Applications PDFs in Downloads were reviewed visually and through text extraction. All are one-page, US Letter documents generated from the same compact LaTeX family."
    )
    resume_table = doc.add_table(rows=1, cols=5)
    resume_table.style = "Table Grid"
    for i, text in enumerate(("Timestamp", "Target lane", "Experience order", "Selected projects", "Review note")):
        resume_table.rows[0].cells[i].text = text
    repeat_table_header(resume_table.rows[0])
    for timestamp, lane, experience, projects, note in RESUME_VARIANTS:
        row = resume_table.add_row().cells
        for i, value in enumerate((timestamp, lane, experience, projects, note)):
            row[i].text = value
    set_table_geometry(resume_table, [1250, 1700, 1800, 2450, 2160])
    for row in resume_table.rows:
        prevent_row_split(row)
    format_table_text(resume_table, size=8.3)

    doc.add_heading("Approved resume layout system", level=2)
    for item in (
        "One US Letter page using Charter, approximately 9 pt body text, compact margins and black section rules.",
        "Header: name plus phone, email, LinkedIn and either GitHub or portfolio; use consistent destinations and display labels.",
        "Section order: Education, Work Experience, job-specific Projects, Technical Skills.",
        "Reorder experience only when role relevance materially improves scanning; do not change dates, titles or evidence strength.",
        "Select two or three projects whose verified evidence matches the role; preserve the same canonical fact set across variants.",
        "Skills should be grouped by practiced capability. Separate hands-on tools from working knowledge and concepts.",
    ):
        doc.add_paragraph(item, style="List Bullet")
    add_label_paragraph(
        doc,
        "Formatting observation",
        "The visual system is clean and ATS-friendly, but the current 9 pt density leaves little room for qualification. Prefer removing unsupported concepts over shrinking text or compressing claims further.",
        fill=LIGHT_BLUE,
    )

    doc.add_heading("Resume claims requiring correction or evidence", level=1)
    doc.add_paragraph(
        "These issues recur across the reviewed PDFs. The tailoring agent should treat each as a prohibited or qualified claim until the required evidence is approved."
    )
    for subject, claim, evidence, action in RESUME_RISKS:
        doc.add_heading(subject, level=2)
        claim_paragraph = add_label_paragraph(doc, "Claim seen in recent resumes", claim)
        claim_paragraph.paragraph_format.keep_with_next = True
        evidence_paragraph = add_label_paragraph(doc, "Evidence assessment", evidence)
        evidence_paragraph.paragraph_format.keep_with_next = True
        add_label_paragraph(doc, "Agent rule", action, fill=PALE_RED)

    doc.add_page_break()
    doc.add_heading("Priority project cross-verification issues", level=1)
    intro = doc.add_paragraph(
        "These are evidence conflicts or missing provenance—not necessarily false claims. Resolve them before automated resume tailoring uses the associated numbers."
    )
    intro.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Project", "Observed conflict", "Required action"]
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
    repeat_table_header(table.rows[0])
    for project, conflict, action in DISCREPANCIES:
        cells = table.add_row().cells
        cells[0].text = project
        cells[1].text = conflict
        cells[2].text = action
    set_table_geometry(table, [1750, 3650, 3960])
    for row in table.rows:
        prevent_row_split(row)
    format_table_text(table, size=8.7)

    doc.add_page_break()
    doc.add_heading("Complete repository inventory", level=1)
    doc.add_paragraph(
        "Language percentages come from GitHub Linguist and may be dominated by generated Unity, report or notebook assets. Commit counts refer to the default branch only."
    )
    inventory = doc.add_table(rows=1, cols=6)
    inventory.style = "Table Grid"
    for i, text in enumerate(("Repository", "Type", "Access", "Primary languages", "Last push", "Resume posture")):
        inventory.rows[0].cells[i].text = text
    repeat_table_header(inventory.rows[0])
    project_by_repo = {p["repo"]: p for p in PROJECTS}
    for meta in metadata.values():
        row = inventory.add_row().cells
        project = project_by_repo.get(meta["name"])
        if meta.get("isFork"):
            kind = "Fork"
        elif meta["name"] in {"tusharpathaknyu", "tusharpathaknyu.github.io"}:
            kind = "Profile"
        else:
            kind = "Original"
        row[0].text = meta["name"]
        row[1].text = kind
        row[2].text = "Private" if meta.get("isPrivate") else "Public"
        row[3].text = language_summary(meta)
        row[4].text = meta.get("pushedAt", "")[:10]
        row[5].text = project["posture"] if project else "Reference only"
    set_table_geometry(inventory, [1740, 800, 800, 2300, 1150, 2570])
    for row in inventory.rows:
        prevent_row_split(row)
    format_table_text(inventory, size=8.4)

    current_category = None
    for project in PROJECTS:
        if project["category"] != current_category:
            if current_category is not None:
                doc.add_page_break()
            current_category = project["category"]
            doc.add_heading(current_category, level=1)
            doc.add_paragraph(
                "Each card separates repository evidence, a conservative resume draft and the specific fact check required before use."
            )
        add_project(doc, project, metadata[project["repo"]])

    doc.add_page_break()
    doc.add_heading("Profile, portfolio and forked repositories", level=1)
    for item in NON_PROJECTS:
        doc.add_heading(item["name"], level=2)
        add_label_paragraph(doc, "Repository type", item["kind"])
        doc.add_paragraph(item["note"])

    doc.add_heading("Recommended data model for the job agent", level=1)
    doc.add_paragraph(
        "After cross-verification, convert each approved project into structured candidate evidence so the tailoring model can select facts without inventing or mixing project states."
    )
    for item in (
        "Project name and canonical repository URL",
        "Approved one-sentence summary and two to four verified accomplishment bullets",
        "Verified technologies, role/ownership, dates and collaboration context",
        "Evidence links or file paths for every numeric claim",
        "Allowed job families (for example: RTL verification, power electronics, ML systems, full-stack)",
        "Prohibited/outdated claims and planned work that must never be presented as completed",
        "Last verification date and reviewer approval status",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Method and limitations", level=1)
    for item in (
        "Sources: five newest Applications PDFs from Downloads; the supplied LaTeX template; candidate-provided Dreamline, Texas Instruments and Xarvis information; private portfolio content; and authenticated GitHub metadata, trees, artifacts and latest Actions runs available on 2026-07-01.",
        "Dreamline/Texas Instruments certificate details were supplied as text in the conversation; the signed original certificate files were not independently inspected in this review.",
        "This review did not clone and execute every repository, inspect external demo links, validate hardware on a bench, or independently reproduce all training/simulation results.",
        "A committed README/report is candidate-authored evidence, not independent verification. Result CSV/JSON artifacts and passing CI increase confidence but do not replace reproducible commands and raw logs.",
        "Private repository content is summarized only in this local dossier. Keep the document private if those projects are not intended for public disclosure.",
        "GitHub ownership and commit attribution do not by themselves prove sole authorship, production use, patents, publications or employer authorization.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Reviewer completion", level=1)
    doc.add_paragraph(
        "Before importing this evidence into the job agent, complete these steps in order:",
    )
    for item in (
        "Confirm contact details, education dates and the scholarship claim against primary records.",
        "Confirm certificate text against signed originals and approve exact work-experience bullets.",
        "Resolve each item in the priority cross-verification table and update the corresponding repository evidence.",
        "Confirm your exact role, collaborators and ownership for every project selected for resume use.",
        "Approve one canonical result table and one resume bullet set per project.",
        "Mark planned, simulated-only and hardware-validated milestones explicitly.",
        "Record the review date and re-run this inventory whenever repository evidence changes materially.",
    ):
        doc.add_paragraph(item, style="List Number")
    signoff = doc.add_paragraph()
    signoff.paragraph_format.space_before = Pt(12)
    signoff.paragraph_format.space_after = Pt(0)
    set_font(signoff.add_run("Reviewer: ______________________________    Date: __________________    Approved for tailoring:  Yes / No"), bold=True, color=DARK_BLUE)

    doc.core_properties.title = "Master Candidate Evidence Dossier — Tushar Dhananjay Pathak"
    doc.core_properties.subject = "Work history, resume patterns and GitHub project evidence for cross-verified tailoring"
    doc.core_properties.author = "Tushar Dhananjay Pathak"
    doc.core_properties.keywords = "resume, work experience, GitHub, projects, candidate evidence, cross-verification"
    doc.core_properties.comments = "Generated from resume, candidate-supplied work evidence and an authenticated GitHub snapshot on 2026-07-01."

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
